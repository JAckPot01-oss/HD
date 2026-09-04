from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile
import json

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def export_csv_zip(tables: dict[str, pd.DataFrame], result: dict | None = None) -> bytes:
    out=BytesIO()
    with ZipFile(out,"w",ZIP_DEFLATED) as z:
        for name,df in tables.items():
            z.writestr(f"{name}.csv",df.to_csv(index=False).encode("utf-8-sig"))
        if result is not None:
            z.writestr("机位计算结果.json",json.dumps(result,ensure_ascii=False,indent=2).encode("utf-8"))
    return out.getvalue()


def _shade(cell, fill: str) -> None:
    tc_pr=cell._tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tc_pr.append(shd)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for edge,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        tag=tcMar.find(qn(f"w:{edge}"))
        if tag is None: tag=OxmlElement(f"w:{edge}"); tcMar.append(tag)
        tag.set(qn("w:w"),str(val)); tag.set(qn("w:type"),"dxa")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table=doc.add_table(rows=1,cols=len(headers))
    table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; table.style="Table Grid"
    widths=widths or [6.5/len(headers)]*len(headers)
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Inches(widths[i]); cell.text=h; _shade(cell,"E8EEF5"); _set_cell_margins(cell)
        for run in cell.paragraphs[0].runs: run.bold=True; run.font.name="Calibri"; run.font.size=Pt(9)
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells=table.add_row().cells
        for i,value in enumerate(row):
            cells[i].width=Inches(widths[i]); cells[i].text=str(value); _set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[i].paragraphs[0].runs: run.font.name="Calibri"; run.font.size=Pt(9)
    return table


def build_calculation_summary_docx(result: dict) -> bytes:
    doc=Document()
    section=doc.sections[0]
    section.page_width=Inches(8.5); section.page_height=Inches(11)
    section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=Inches(1)
    section.header_distance=section.footer_distance=Inches(0.492)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for name,size,before,after,color in [("Heading 1",16,16,8,"2E74B5"),("Heading 2",13,12,6,"2E74B5"),("Heading 3",12,8,4,"1F4D78")]:
        s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=section.header.paragraphs[0]; header.text="液压爬模设计辅助系统 · 计算摘要"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor(100,110,120)
    title=doc.add_paragraph(); title.paragraph_format.space_after=Pt(4)
    run=title.add_run("液压爬模机位计算摘要"); run.bold=True; run.font.name="Calibri"; run.font.size=Pt(23)
    subtitle=doc.add_paragraph(f"{result['inputs']['project_name']} · {result['inputs']['position_name']} · {result['model']}")
    subtitle.paragraph_format.space_after=Pt(14); subtitle.runs[0].font.size=Pt(12); subtitle.runs[0].font.color.rgb=RGBColor(80,80,80)
    doc.add_heading("结论",level=1)
    passed=all(result["checks"].values())
    p=doc.add_paragraph(); r=p.add_run("本次自动验算结果："+("全部通过" if passed else "存在未通过项目")); r.bold=True; r.font.color.rgb=RGBColor(31,90,60) if passed else RGBColor(155,28,28)
    doc.add_paragraph("本文件为一期原型生成的计算摘要，正式工程成果仍须由具备资格的设计人员复核、签署并按现行标准审查。")
    doc.add_heading("关键输入",level=1)
    inp=result["inputs"]
    input_rows=[["机位承担几何","机位间距/悬挑",f"{inp['spacing_large_mm']:.0f} / {inp['spacing_small_mm']:.0f} / {inp['cantilever_mm']:.0f} mm"],
                ["结构","层高/墙厚",f"{inp['floor_height_mm']:.0f} / {inp['wall_thickness_mm']:.0f} mm"],
                ["模板","高度/重量",f"{inp['template_height_mm']:.0f} mm / {inp['template_weight_kg_m2']:.1f} kg/m²"],
                ["环境","阵风/高度变化系数",f"{inp['gust_factor']:.3f} / {inp['height_factor']:.3f}"],
                ["连接","附墙螺栓数量",str(inp['bolt_count'])]]
    _add_table(doc,["类别","参数","数值"],input_rows,[1.1,2.2,3.2])
    doc.add_heading("荷载与内力",level=1)
    rows=[]
    for state in ["施工工况","爬升工况","停工工况"]:
        rows.append([state,f"{result['tension_total_kn'][state]:.3f}",f"{result['shear_total_kn'][state]:.3f}",f"{result['bolt_interaction'][state]:.3f}","满足" if result['bolt_interaction'][state] <= 1 else "不满足"])
    _add_table(doc,["工况","总拉力(kN)","总剪力(kN)","螺栓组合比","判定"],rows,[1.3,1.35,1.35,1.35,1.15])
    doc.add_heading("构件验算",level=1)
    check_rows=[[name,"满足" if ok else "不满足"] for name,ok in result["checks"].items()]
    _add_table(doc,["验算项目","结果"],check_rows,[4.8,1.7])
    doc.add_heading("主要结果",level=1)
    key_rows=[
        ["爬模装置自重标准值",f"{result['self_weight_kn']:.3f} kN"],
        ["混凝土冲切承载力",f"{result['concrete']['冲切承载力(kN)']:.3f} kN"],
        ["混凝土局部受压承载力",f"{result['concrete']['局部受压承载力(kN)']:.3f} kN"],
        ["导轨跨中挠度",f"{result['rail']['挠度(mm)']:.3f} mm（允许 {result['rail']['允许挠度(mm)']:.3f} mm）"],
    ]
    _add_table(doc,["指标","结果"],key_rows,[3.1,3.4])
    out=BytesIO(); doc.save(out); return out.getvalue()
